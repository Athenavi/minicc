"""Word 文档生成 — python-docx 驱动。"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field

from app.models.permission import PermissionLevel
from app.models.tool import ToolResult
from app.tools.base import BaseTool, ToolCategory, ToolUseContext


class WordCreateInput(BaseModel):
    path: str = Field(description="Output file path")
    title: Optional[str] = Field(default=None, description="Document title")
    content: str = Field(description="Document content (supports basic Markdown-like formatting)")


class WordCreateTool(BaseTool):
    name = "word_create"
    description = "Create a Word document with title and content."
    input_schema = WordCreateInput
    permission_level = PermissionLevel.WRITE
    category = ToolCategory.FILE

    async def execute(self, input_data: WordCreateInput, context: ToolUseContext | None = None) -> ToolResult:
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            return ToolResult(tool_call_id="", output="[word] python-docx not installed. Run: pip install python-docx", is_error=True)

        path = Path(input_data.path).expanduser().resolve()
        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        if input_data.title:
            doc.add_heading(input_data.title, level=1)

        for para_text in input_data.content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("## "):
                doc.add_heading(para_text[3:], level=2)
            elif para_text.startswith("# "):
                doc.add_heading(para_text[2:], level=1)
            elif para_text.startswith("- ") or para_text.startswith("* "):
                for line in para_text.split("\n"):
                    line = line.strip()
                    if line.startswith("- ") or line.startswith("* "):
                        doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(para_text)

        doc.save(path)
        return ToolResult(tool_call_id="", output=f"[word] Created: {path.name} ({len(input_data.content)} chars)")


class WordTemplateInput(BaseModel):
    template_path: str = Field(description="Path to .docx template file")
    output_path: str = Field(description="Output file path")
    replacements: dict[str, str] = Field(description="Map of {{placeholder}} → value")


class WordTemplateTool(BaseTool):
    name = "word_template"
    description = "Fill a Word template by replacing placeholders with values."
    input_schema = WordTemplateInput
    permission_level = PermissionLevel.WRITE
    category = ToolCategory.FILE

    async def execute(self, input_data: WordTemplateInput, context: ToolUseContext | None = None) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(tool_call_id="", output="[word] python-docx not installed.", is_error=True)

        tpl = Path(input_data.template_path).expanduser()
        if not tpl.exists():
            return ToolResult(tool_call_id="", output=f"[word] Template not found: {tpl}", is_error=True)

        doc = Document(tpl)
        for para in doc.paragraphs:
            for key, val in input_data.replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, val)

        out = Path(input_data.output_path).expanduser()
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out)
        return ToolResult(tool_call_id="", output=f"[word] Template filled: {out.name}")


def register_word_tools(registry) -> None:
    registry.register(WordCreateTool())
    registry.register(WordTemplateTool())
